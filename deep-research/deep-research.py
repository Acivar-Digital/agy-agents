import argparse
import json
import os
import sys
import time
import urllib.error
import urllib.request
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
import markdown
from bs4 import BeautifulSoup, Tag
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from weasyprint import HTML, CSS

# Resolve repository root
RESEARCH_DIR = Path(__file__).resolve().parent
REPO_ROOT = RESEARCH_DIR.parent
load_dotenv(dotenv_path=REPO_ROOT / ".env")

PROMPTS_DIR = RESEARCH_DIR / "prompts"
REPORTS_DIR = RESEARCH_DIR / "reports"

# Ensure directories exist
PROMPTS_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

LITEROUTER_PORT = os.getenv("LITEROUTER_PORT", "7766")
LITEROUTER_KEY = os.getenv("LITEROUTER_AUTH_KEY", "YOUR_KEY_HERE")
GATEWAY_URL = f"http://localhost:{LITEROUTER_PORT}/v1beta/interactions"

CITI_PDF_CSS = """
@page {
    size: A4;
    margin: 20mm 15mm 20mm 15mm;
    @top-right {
        content: "CITI INSTITUTIONAL RESEARCH";
        font-family: Arial, sans-serif;
        font-size: 8pt;
        font-weight: bold;
        color: #002D62;
    }
    @bottom-left {
        content: "CONFIDENTIAL - FOR INSTITUTIONAL CLIENT USE ONLY";
        font-family: Arial, sans-serif;
        font-size: 7.5pt;
        color: #718096;
    }
    @bottom-right {
        content: "Page " counter(page) " of " counter(pages);
        font-family: Arial, sans-serif;
        font-size: 8pt;
        color: #002D62;
    }
}

body {
    font-family: 'Helvetica Neue', Arial, sans-serif;
    color: #2D3748 !important;
    background-color: #FFFFFF !important;
    line-height: 1.6;
    font-size: 9.5pt;
}

h1 {
    color: #002D62 !important;
    border-bottom: 2.5px solid #D9261C;
    padding-bottom: 6px;
    font-size: 18pt;
    margin-top: 0;
    margin-bottom: 12px;
}

h2 {
    color: #002D62 !important;
    font-size: 13pt;
    border-bottom: 1px solid #CBD5E0;
    padding-bottom: 4px;
    margin-top: 18px;
    margin-bottom: 8px;
}

h3 {
    color: #D9261C !important;
    font-size: 11pt;
    margin-top: 14px;
    margin-bottom: 6px;
}

blockquote {
    background-color: #F7FAFC !important;
    border-left: 4px solid #002D62 !important;
    margin: 14px 0;
    padding: 10px 14px;
    font-size: 9.5pt;
    color: #002D62 !important;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 14px 0;
    background-color: #FFFFFF !important;
    color: #2D3748 !important;
    font-size: 8.5pt;
}

th {
    background-color: #002D62 !important;
    color: #FFFFFF !important;
    text-align: left;
    padding: 8px 10px;
    font-weight: bold;
    border: 1px solid #002D62 !important;
}

td {
    padding: 8px 10px;
    background-color: #FFFFFF !important;
    color: #2D3748 !important;
    border: 1px solid #E2E8F0 !important;
}

tr:nth-child(even) td {
    background-color: #F8FAFC !important;
}

code {
    font-family: 'Courier New', Courier, monospace;
    background-color: #EDF2F7 !important;
    color: #2D3748 !important;
    padding: 2px 4px;
    border-radius: 3px;
    font-size: 8.5pt;
}

pre {
    background-color: #1A202C !important;
    color: #F7FAFC !important;
    padding: 10px;
    border-radius: 4px;
    overflow-x: auto;
}

details {
    margin: 10px 0;
    padding: 8px;
    background-color: #EDF2F7 !important;
    border-radius: 4px;
}

summary {
    font-weight: bold;
    color: #002D62 !important;
    cursor: pointer;
}
"""


def export_to_html(md_text: str, output_html_path: Path):
    """Converts Markdown text into a clean styled HTML web report ready for PDF printing."""
    try:
        html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code", "toc"])
        full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Institutional Deep Research Report</title>
    <style>
        :root {{
            --citi-navy: #002D62;
            --citi-red: #D9261C;
            --bg-card: #F8FAFC;
            --text-main: #1E293B;
            --text-muted: #64748B;
            --border-color: #E2E8F0;
        }}
        body {{
            font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
            background-color: #F8FAFC;
            color: var(--text-main);
            margin: 0;
            padding: 20px;
            line-height: 1.6;
            -webkit-print-color-adjust: exact;
            print-color-adjust: exact;
        }}
        .container {{
            max-width: 1100px;
            margin: 20px auto;
            background: #FFFFFF;
            padding: 50px 60px;
            border-radius: 8px;
            box-shadow: 0 10px 25px rgba(0, 0, 0, 0.05);
            border: 1px solid var(--border-color);
        }}
        h1 {{
            color: var(--citi-navy);
            font-size: 26px;
            font-weight: 700;
            border-bottom: 3px solid var(--citi-red);
            padding-bottom: 10px;
            margin-top: 0;
            margin-bottom: 24px;
        }}
        h2 {{
            color: var(--citi-navy);
            font-size: 20px;
            font-weight: 700;
            border-bottom: 2px solid var(--border-color);
            padding-bottom: 8px;
            margin-top: 36px;
            margin-bottom: 16px;
        }}
        h3 {{
            color: var(--citi-red);
            font-size: 16px;
            font-weight: 700;
            margin-top: 24px;
            margin-bottom: 12px;
        }}
        blockquote {{
            background-color: var(--bg-card);
            border-left: 5px solid var(--citi-navy);
            margin: 24px 0;
            padding: 18px 24px;
            border-radius: 0 6px 6px 0;
            color: var(--citi-navy);
            font-size: 15px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 24px 0;
            background-color: #FFFFFF !important;
            font-size: 14px;
            box-shadow: 0 2px 6px rgba(0,0,0,0.02);
            border-radius: 6px;
            overflow: hidden;
        }}
        th {{
            background-color: var(--citi-navy) !important;
            color: #FFFFFF !important;
            text-align: left;
            padding: 12px 16px;
            font-weight: 600;
            border: 1px solid var(--citi-navy) !important;
        }}
        td {{
            padding: 12px 16px;
            background-color: #FFFFFF !important;
            color: var(--text-main) !important;
            border: 1px solid var(--border-color) !important;
        }}
        tr:nth-child(even) td {{
            background-color: #F8FAFC !important;
        }}
        tr:hover td {{
            background-color: #F1F5F9 !important;
        }}
        code {{
            font-family: 'Consolas', 'Courier New', monospace;
            background-color: #F1F5F9 !important;
            color: #0F172A !important;
            padding: 3px 6px;
            border-radius: 4px;
            font-size: 13px;
            border: 1px solid #E2E8F0;
        }}
        pre {{
            background-color: #F8FAFC !important;
            color: #0F172A !important;
            padding: 16px;
            border-radius: 6px;
            border: 1px solid var(--border-color);
            overflow-x: auto;
        }}
        details {{
            background-color: #F8FAFC;
            border: 1px solid var(--border-color);
            padding: 14px 18px;
            border-radius: 6px;
            margin: 20px 0;
        }}
        summary {{
            font-weight: 700;
            color: var(--citi-navy);
            cursor: pointer;
        }}
        @media print {{
            body {{
                background-color: #FFFFFF;
                padding: 0;
            }}
            .container {{
                box-shadow: none;
                border: none;
                max-width: 100%;
                padding: 0;
                margin: 0;
            }}
            th {{
                background-color: var(--citi-navy) !important;
                color: #FFFFFF !important;
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }}
            blockquote {{
                background-color: var(--bg-card) !important;
                border-left: 5px solid var(--citi-navy) !important;
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }}
            tr:nth-child(even) td {{
                background-color: #F8FAFC !important;
                -webkit-print-color-adjust: exact;
                print-color-adjust: exact;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        {html_body}
    </div>
</body>
</html>"""
        output_html_path.write_text(full_html, encoding="utf-8")
        print(f"🌐 Clean HTML Report generated: {output_html_path}")
    except Exception as e:
        print(f"⚠️ Failed to generate HTML: {e}")


def export_to_pdf(md_text: str, output_pdf_path: Path):
    """Converts Markdown text into a styled Citi PDF report via WeasyPrint."""
    try:
        html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code", "toc"])
        full_html = f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>{html_body}</body></html>"
        HTML(string=full_html).write_pdf(output_pdf_path, stylesheets=[CSS(string=CITI_PDF_CSS)])
        print(f"📄 Citi PDF Report generated:  {output_pdf_path}")
    except Exception as e:
        print(f"⚠️ Failed to generate PDF: {e}")


def _set_cell_background(cell, fill_hex: str):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def export_to_docx(md_text: str, output_docx_path: Path):
    """Converts Markdown text into a formatted Citi Word document via python-docx."""
    try:
        doc = docx.Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("CITI INSTITUTIONAL RESEARCH")
        hrun.font.name = "Calibri"
        hrun.font.bold = True
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0, 45, 98)

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        frun = fp.add_run("CONFIDENTIAL - FOR INSTITUTIONAL CLIENT USE ONLY")
        frun.font.name = "Calibri"
        frun.font.size = Pt(8)
        frun.font.color.rgb = RGBColor(113, 128, 150)

        html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
        soup = BeautifulSoup(html_body, "html.parser")

        for elem in soup.children:
            if not isinstance(elem, Tag):
                continue
            if elem.name == "h1":
                h = doc.add_heading(elem.get_text(), level=1)
                for r in h.runs:
                    r.font.name = "Calibri"
                    r.font.color.rgb = RGBColor(0, 45, 98)
                    r.font.bold = True
            elif elem.name == "h2":
                h = doc.add_heading(elem.get_text(), level=2)
                for r in h.runs:
                    r.font.name = "Calibri"
                    r.font.color.rgb = RGBColor(0, 45, 98)
                    r.font.bold = True
            elif elem.name == "h3":
                h = doc.add_heading(elem.get_text(), level=3)
                for r in h.runs:
                    r.font.name = "Calibri"
                    r.font.color.rgb = RGBColor(217, 38, 28)
                    r.font.bold = True
            elif elem.name in ["h4", "h5", "h6"]:
                h = doc.add_heading(elem.get_text(), level=4)
                for r in h.runs:
                    r.font.name = "Calibri"
                    r.font.color.rgb = RGBColor(0, 45, 98)
            elif elem.name == "blockquote":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                r = p.add_run(elem.get_text().strip())
                r.font.name = "Calibri"
                r.font.italic = True
                r.font.color.rgb = RGBColor(0, 45, 98)
            elif elem.name == "table":
                rows = [r for r in elem.find_all("tr") if isinstance(r, Tag)]
                if rows:
                    headers = [th.get_text().strip() for th in rows[0].find_all(["th", "td"])]
                    if headers:
                        table = doc.add_table(rows=len(rows), cols=len(headers))
                        table.style = "Table Grid"
                        for col_idx, htext in enumerate(headers):
                            cell = table.cell(0, col_idx)
                            cell.text = htext
                            _set_cell_background(cell, "002D62")
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.font.name = "Calibri"
                                    r.font.color.rgb = RGBColor(255, 255, 255)
                                    r.font.bold = True
                                    r.font.size = Pt(9)
                        for row_idx, row in enumerate(rows[1:], 1):
                            cells = [td.get_text().strip() for td in row.find_all("td")]
                            bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
                            for col_idx, ctext in enumerate(cells):
                                if col_idx < len(headers):
                                    cell = table.cell(row_idx, col_idx)
                                    cell.text = ctext
                                    if bg_color != "FFFFFF":
                                        _set_cell_background(cell, bg_color)
                                    for p in cell.paragraphs:
                                        for r in p.runs:
                                            r.font.name = "Calibri"
                                            r.font.size = Pt(9)
                                            r.font.color.rgb = RGBColor(45, 55, 72)
            elif elem.name in ["ul", "ol"]:
                for li in [item for item in elem.find_all("li") if isinstance(item, Tag)]:
                    p = doc.add_paragraph(style="List Bullet" if elem.name == "ul" else "List Number")
                    r = p.add_run(li.get_text().strip())
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)
            elif elem.name == "p":
                p = doc.add_paragraph()
                r = p.add_run(elem.get_text().strip())
                r.font.name = "Calibri"
                r.font.size = Pt(10)

        doc.save(output_docx_path)
        print(f"📝 Citi DOCX Report generated: {output_docx_path}")
    except Exception as e:
        print(f"⚠️ Failed to generate DOCX: {e}")


def transform_json_to_report(res_json: dict, elapsed: float | None = None) -> str:
    """Transforms raw Antigravity interaction JSON response into a rich readable Markdown report."""
    output_text_parts = []

    # 1. Check top-level keys first
    for k in ["output_text", "output", "response", "answer", "result"]:
        v = res_json.get(k)
        if isinstance(v, str) and v.strip():
            output_text_parts.append(v.strip())
        elif isinstance(v, dict) and (v.get("text") or v.get("message")):
            output_text_parts.append(str(v.get("text") or v.get("message")).strip())

    # 2. Parse execution steps
    steps = res_json.get("steps", [])
    thoughts = []
    search_queries = []
    sources = []
    sources_seen = set()

    for step in steps:
        stype = step.get("type")
        if stype == "model_output":
            contents = step.get("content", [])
            for c in contents:
                if isinstance(c, dict) and c.get("text"):
                    output_text_parts.append(c["text"].strip())
                elif isinstance(c, str) and c.strip():
                    output_text_parts.append(c.strip())
        elif stype == "thought":
            summaries = step.get("summary", [])
            if isinstance(summaries, list):
                for s in summaries:
                    if isinstance(s, dict) and s.get("text"):
                        thoughts.append(s["text"].strip())
                    elif isinstance(s, str) and s.strip():
                        thoughts.append(s.strip())
            elif isinstance(summaries, str) and summaries.strip():
                thoughts.append(summaries.strip())
        elif stype == "google_search_call":
            args = step.get("arguments", {})
            if isinstance(args, dict):
                queries = args.get("queries", [])
                for q in queries:
                    if q not in search_queries:
                        search_queries.append(q)
        elif stype == "google_search_result":
            results = step.get("result", [])
            for r in results:
                s_sug = r.get("search_suggestions")
                if s_sug:
                    try:
                        parsed = json.loads(s_sug)
                        for f in parsed.get("fields", []):
                            if f.get("name") == "result":
                                vals = f.get("value", {}).get("listValue", {}).get("values", [])
                                for v in vals:
                                    q_struct = v.get("structValue", {}).get("fields", [])
                                    for qf in q_struct:
                                        if qf.get("name") == "results":
                                            res_items = qf.get("value", {}).get("listValue", {}).get("values", [])
                                            for res_item in res_items:
                                                item_fields = res_item.get("structValue", {}).get("fields", [])
                                                item_dict = {}
                                                for ifield in item_fields:
                                                    val_obj = ifield.get("value", {})
                                                    if "stringValue" in val_obj:
                                                        item_dict[ifield["name"]] = val_obj["stringValue"]
                                                title = item_dict.get("source_title", "Untitled Source")
                                                url = item_dict.get("url")
                                                snippet = item_dict.get("snippet", "")
                                                pub_time = item_dict.get("publication_time", "")
                                                if url and url not in sources_seen:
                                                    sources_seen.add(url)
                                                    sources.append({
                                                        "title": title,
                                                        "url": url,
                                                        "snippet": snippet,
                                                        "date": pub_time,
                                                    })
                    except Exception:
                        pass

    full_output_text = "\n\n".join(output_text_parts).strip()

    report_lines = []
    agent_name = res_json.get("agent", "antigravity-preview-05-2026")
    env_id = res_json.get("environment_id", "N/A")
    usage = res_json.get("usage", {})

    if full_output_text:
        report_lines.append(full_output_text)
        report_lines.append("\n\n---\n\n")
    else:
        report_lines.append("# Deep Research Report\n\n*(No model output text found in JSON response)*\n\n---\n\n")

    report_lines.append("## Appendix: Agent Execution & Grounding Provenance\n\n")

    report_lines.append("### Execution Metadata\n")
    report_lines.append(f"- **Agent:** `{agent_name}`\n")
    report_lines.append(f"- **Environment ID:** `{env_id}`\n")
    if elapsed:
        report_lines.append(f"- **Execution Time:** {elapsed:.1f} seconds\n")
    if usage:
        tot_tok = usage.get("total_tokens", 0)
        in_tok = usage.get("total_input_tokens", 0)
        out_tok = usage.get("total_output_tokens", 0)
        thought_tok = usage.get("total_thought_tokens", 0)
        report_lines.append(
            f"- **Token Usage:** {tot_tok:,} total ({in_tok:,} input, {out_tok:,} output, {thought_tok:,} thought)\n"
        )
    report_lines.append("\n")

    if search_queries:
        report_lines.append("### Verified Search Queries\n")
        for q in search_queries:
            report_lines.append(f"- `{q}`\n")
        report_lines.append("\n")

    if sources:
        report_lines.append("### Grounding Sources & References\n")
        for i, src in enumerate(sources, 1):
            stitle = src.get("title", "Source")
            surl = src.get("url", "#")
            ssnip = src.get("snippet", "")
            d = src.get("date", "")
            date_str = f" ({d})" if d else ""
            report_lines.append(f"{i}. [{stitle}]({surl}){date_str}\n")
            if ssnip:
                report_lines.append(f"   > *{ssnip}*\n")
        report_lines.append("\n")

    if thoughts:
        report_lines.append("<details>\n<summary><b>Agent Reasoning & Thought Process Log (Expand)</b></summary>\n\n")
        for t in thoughts:
            report_lines.append(f"{t}\n\n")
        report_lines.append("</details>\n")

    return "".join(report_lines)


def run_deep_research():
    parser = argparse.ArgumentParser(description="Deep Research Agent Tool")
    parser.add_argument("target", help="The name of the prompt (e.g. Direction_of_JPY) or path to JSON if --transform")
    parser.add_argument("--transform", action="store_true", help="Transform an existing raw JSON response into markdown, pdf, and docx reports without making an API call")
    args = parser.parse_args()

    if args.transform:
        # User passed a path to a JSON file
        raw_json_file = Path(args.target)
        if not raw_json_file.exists():
            print(f"❌ Error: Raw JSON file not found at {raw_json_file}")
            sys.exit(1)
        
        # Derive output file paths
        if raw_json_file.name.endswith("_raw.json"):
            output_md = raw_json_file.with_name(raw_json_file.name.replace("_raw.json", ".md"))
            output_html = raw_json_file.with_name(raw_json_file.name.replace("_raw.json", ".html"))
            output_pdf = raw_json_file.with_name(raw_json_file.name.replace("_raw.json", ".pdf"))
            output_docx = raw_json_file.with_name(raw_json_file.name.replace("_raw.json", ".docx"))
        else:
            output_md = raw_json_file.with_suffix(".md")
            output_html = raw_json_file.with_suffix(".html")
            output_pdf = raw_json_file.with_suffix(".pdf")
            output_docx = raw_json_file.with_suffix(".docx")

        print(f"🔄 Transforming existing {raw_json_file} into .md, .html, .pdf, and .docx...")
        res_json = json.loads(raw_json_file.read_text(encoding="utf-8"))
        final_report = transform_json_to_report(res_json)
        output_md.write_text(final_report, encoding="utf-8")
        print(f"🎉 Successfully transformed Markdown report: {output_md}")
        
        # Multi-format exports
        export_to_html(final_report, output_html)
        export_to_pdf(final_report, output_pdf)
        export_to_docx(final_report, output_docx)
        sys.exit(0)

    # Normal execution mode
    prompt_stem = Path(args.target).stem
    prompt_file = PROMPTS_DIR / f"{prompt_stem}.md"

    if not prompt_file.exists():
        print(f"❌ Error: Prompt template not found at {prompt_file}")
        sys.exit(1)

    # Setup output files
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_md = REPORTS_DIR / f"{prompt_stem}_{timestamp}.md"
    output_html = REPORTS_DIR / f"{prompt_stem}_{timestamp}.html"
    output_pdf = REPORTS_DIR / f"{prompt_stem}_{timestamp}.pdf"
    output_docx = REPORTS_DIR / f"{prompt_stem}_{timestamp}.docx"
    output_json = REPORTS_DIR / f"{prompt_stem}_{timestamp}_raw.json"

    prompt_content = prompt_file.read_text(encoding="utf-8").strip()
    print("==================================================================")
    print("🔬 DEEP RESEARCH AGENT (via LiteRouter + Antigravity)")
    print("==================================================================")
    print(f"📍 Target Gateway: {GATEWAY_URL}")
    print(f"📄 Reading Prompt: {prompt_file}")
    print(f"📁 Output Dir:     {REPORTS_DIR}")
    print("------------------------------------------------------------------")
    print(prompt_content[:300] + ("..." if len(prompt_content) > 300 else ""))
    print("==================================================================\n")

    payload = {
        "agent": "antigravity-preview-05-2026",
        "input": prompt_content,
        "environment": "remote",
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {LITEROUTER_KEY}",
    }

    req = urllib.request.Request(
        GATEWAY_URL,
        data=json.dumps(payload).encode("utf-8"),
        headers=headers,
        method="POST",
    )

    print("🚀 Dispatching request to Antigravity Agent (Google Cloud Remote Sandbox)...")
    print("⏳ This multi-step execution takes approximately 1-3 minutes. Please wait...\n")

    start_time = time.time()
    try:
        with urllib.request.urlopen(req, timeout=300) as resp:
            elapsed = time.time() - start_time
            print(f"✅ Response received in {elapsed:.1f} seconds (HTTP {resp.status})!")

            res_body = resp.read().decode("utf-8")
            res_json = json.loads(res_body)

            # --- 1. SAVE RAW JSON FOR DEBUGGING ---
            output_json.write_text(json.dumps(res_json, indent=2), encoding="utf-8")
            print(f"💾 Saved raw API response to: {output_json}")

            # --- 2. TRANSFORM JSON INTO READABLE REPORT ---
            final_report_str = transform_json_to_report(res_json, elapsed=elapsed)
            output_md.write_text(final_report_str, encoding="utf-8")

            print("\n🎉 Deep Research Complete!")
            print(f"📄 Markdown report: {output_md}")
            
            # --- 3. EXPORT TO HTML, PDF & DOCX ---
            export_to_html(final_report_str, output_html)
            export_to_pdf(final_report_str, output_pdf)
            export_to_docx(final_report_str, output_docx)
            
            print(f"📊 Usage stats: {json.dumps(res_json.get('usage', {}), indent=2)}")
            sys.exit(0)

    except urllib.error.HTTPError as e:
        err_body = e.read().decode("utf-8", errors="ignore")
        print(f"❌ HTTP Error {e.code}: {err_body[:500]}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Execution Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    run_deep_research()
